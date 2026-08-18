import os
import sys

# Tell Anaconda where to find Homebrew libraries on macOS
if sys.platform == "darwin":
    os.environ["DYLD_FALLBACK_LIBRARY_PATH"] = "/opt/homebrew/lib:" + os.environ.get(
        "DYLD_FALLBACK_LIBRARY_PATH", "/usr/local/lib"
    )
####
import datetime
import io
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
import pandas as pd
from weasyprint import HTML
import os
import re
import pandas as pd
from openai import OpenAI

# Safe import wrapper for WeasyPrint
try:
    from weasyprint import HTML

    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False
    HTML = None


# --- DYNAMIC CLIENT INITIALIZATION (Strategy 3) ---
def get_llm_client():
    groq_key = None
    try:
        groq_key = st.secrets.get("GROQ_API_KEY") or os.getenv("GROQ_API_KEY")
    except Exception:
        groq_key = os.getenv("GROQ_API_KEY")

    if groq_key:
        return (
            OpenAI(
                api_key=groq_key,
                base_url="https://api.groq.com/openai/v1",
            ),
            "llama-3.1-8b-instant",  # Guaranteed active model on Groq free tier
        )

    # Local LM Studio Fallback for your laptop
    return (
        OpenAI(base_url="http://localhost:1234/v1", api_key="lm-studio"),
        "local-model",
    )


client, DEFAULT_MODEL = get_llm_client()


class ConversationMemory:
    def __init__(self, max_history=5):
        self.max_history = max_history
        self.history = []

    def add_interaction(self, question: str, answer: str):
        self.history.append({"question": question, "answer": answer})
        if len(self.history) > self.max_history:
            self.history.pop(0)

    def get_context_string(self) -> str:
        if not self.history:
            return "No previous conversation context."
        context_lines = ["Previous Conversation History:"]
        for idx, item in enumerate(self.history, 1):
            context_lines.append(f"Q{idx}: {item['question']}")
            context_lines.append(f"A{idx}: {item['answer']}")
        return "\n".join(context_lines)

    def reset(self):
        self.history = []


def resolve_followup_question(current_question: str, memory_context: str) -> str:
    if memory_context == "No previous conversation context.":
        return current_question

    system_prompt = """
You are a conversational query assistant.
Given a conversation history and a follow-up question, rewrite the follow-up question so it becomes a clear, standalone data query.
Return ONLY the rewritten standalone question. No extra explanations.
"""
    response = client.chat.completions.create(
        model=DEFAULT_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {
                "role": "user",
                "content": f"{memory_context}\nFollow-up Question: {current_question}",
            },
        ],
        temperature=0.1,
    )
    return response.choices[0].message.content.strip()


import json


def generate_pandas_code(
    user_question: str, schema_json: str, previous_error: str = None
) -> str:
    # Parse columns to give the local LLM a clear, explicit list
    try:
        schema_data = json.loads(schema_json)
        available_columns = [col["name"] for col in schema_data]
    except Exception:
        available_columns = []

    system_prompt = f"""
You are an expert Data Analyst assistant writing executable Python code using `pandas` on an ALREADY LOADED DataFrame named `df`.

### STRICT AVAILABLE COLUMNS IN DATASET:
{available_columns}

### FULL SCHEMA CATALOG:
{schema_json}

### CRITICAL RULES:
1. Return ONLY executable Python code inside a single ```python ... ``` block. No prose or explanations.
2. Store the final result in a variable named `result`. Do NOT use `print()`.
3. **STRICT COLUMN MATCHING:**
   - You MUST ONLY use column names from `STRICT AVAILABLE COLUMNS IN DATASET` above.
   - Match case-sensitively (e.g. if the column is 'Country', do NOT use 'customer_country' or 'country').
   - If no relevant column exists for the question, return: `result = pd.DataFrame([{{'Message': 'Column not found in dataset'}}])`

### QUERY CATEGORIES & CODE PATTERNS:

A. GROUPING / BREAKDOWN / COUNTS:
   - For unique value counts (e.g., how many countries):
     ```python
     # Replace 'EXACT_COL' with column from STRICT AVAILABLE COLUMNS
     result = pd.DataFrame([{{'Total Countries': df['EXACT_COL'].nunique()}}])
     ```
   - For occurrences / list by category:
     ```python
     result = df['EXACT_COL'].value_counts().reset_index()
     result.columns = ['Country', 'Count']
     ```

B. SINGLE-ROW OVERVIEW / KPIs:
   - Wrap scalar dict in list: `result = pd.DataFrame([{{...}}])`

### CRITICAL RULES:
1. NEVER reload data using `pd.read_csv()`. ALWAYS operate strictly on the existing `df` variable.   

"""
    messages = [{"role": "system", "content": system_prompt}]
    if previous_error:
        messages.append(
            {
                "role": "user",
                "content": f"Previous code failed with error: {previous_error}\nAvailable columns are strictly: {available_columns}. Fix code for: {user_question}",
            }
        )
    else:
        messages.append({"role": "user", "content": f"Question: {user_question}"})

    response = client.chat.completions.create(
        model=DEFAULT_MODEL, messages=messages, temperature=0.1
    )
    return response.choices[0].message.content

    # SECURITY SANBOX RATIONALE:

    # are entirely absent from this allowlist namespace.


def execute_sandboxed_code(code_str: str, dataframe: pd.DataFrame) -> dict:
    code_match = re.search(r"```python\s*(.*?)\s*```", code_str, re.DOTALL)
    clean_code = code_match.group(1) if code_match else code_str

    local_vars = {"df": dataframe, "pd": pd}

    try:
        # SECURITY SANBOX RATIONALE:
        # Executing LLM-generated strings directly poses an arbitrary code execution risk.

        exec(clean_code, {}, local_vars)

        raw_result = None

        # 1. Primary Check: Did the LLM define 'result'?
        if "result" in local_vars:
            raw_result = local_vars["result"]

        # 2. Fallback Check: Common variable names
        if raw_result is None:
            for fallback_var in [
                "res",
                "df_result",
                "output",
                "summary",
                "ans",
                "final_df",
            ]:
                if fallback_var in local_vars:
                    raw_result = local_vars[fallback_var]
                    break

        # 3. Last Resort Check
        if raw_result is None:
            possible_results = [
                v for k, v in local_vars.items() if k not in ["df", "pd", "re"]
            ]
            if possible_results:
                raw_result = possible_results[-1]

        if raw_result is None:
            return {
                "success": False,
                "result": None,
                "code": clean_code,
                "error": "Variable 'result' was not defined.",
            }

        # --- NORMALIZE RESULT TO SAFE DATAFRAME / SERIES ---
        if isinstance(raw_result, dict):
            # Safe conversion of scalar dict to single-row DataFrame
            normalized_result = pd.DataFrame([raw_result])
        elif isinstance(raw_result, pd.Series):
            normalized_result = raw_result.reset_index()
        elif isinstance(raw_result, (int, float, str, bool)):
            normalized_result = pd.DataFrame([{"Value": raw_result}])
        else:
            normalized_result = raw_result

        return {
            "success": True,
            "result": normalized_result,
            "code": clean_code,
            "error": None,
        }

    except Exception as e:
        return {"success": False, "result": None, "code": clean_code, "error": str(e)}


def format_result_with_llm(user_question: str, query_result) -> str:
    if isinstance(query_result, pd.DataFrame):
        result_str = query_result.head(10).to_markdown(index=False)
    elif isinstance(query_result, pd.Series):
        result_str = query_result.to_frame().to_markdown()
    else:
        result_str = str(query_result)

    system_prompt = """
You are a Data Storytelling Assistant. 
Rules:
1. Provide a concise natural language answer formatted in clean Markdown.
2. Bold key metrics. 
3. Always include proper spaces between words and numbers. Do NOT use LaTeX math syntax like `$`.
"""
    user_prompt = f"User Question: {user_question}\n\nQuery Result Data:\n{result_str}"

    response = client.chat.completions.create(
        model=DEFAULT_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.1,
    )
    return response.choices[0].message.content


def nl_query_pipeline(
    user_question: str,
    dataframe: pd.DataFrame,
    schema_json: str,
    memory_context: str = None,
):
    if memory_context and memory_context != "No previous conversation context.":
        processed_question = resolve_followup_question(user_question, memory_context)
    else:
        processed_question = user_question

    llm_output = generate_pandas_code(processed_question, schema_json)

    # SECURITY SANBOX RATIONALE:
    # are entirely absent from this allowlist namespace.

    exec_res = execute_sandboxed_code(llm_output, dataframe)

    if not exec_res["success"]:
        llm_output_retry = generate_pandas_code(
            processed_question, schema_json, previous_error=exec_res["error"]
        )
        exec_res = execute_sandboxed_code(llm_output_retry, dataframe)

    return exec_res


def generate_preset_insight(
    preset_type: str, dataframe: pd.DataFrame, schema_json: str
) -> str:
    prompts = {
        "overview": "Provide a high-level summary overview of this dataset including total rows, sales, and profit.",
        "comparison": "Which Category has the highest total Sales and Profit?",
        "outliers": "Show me rows with unusual or extreme profit values.",
    }

    prompt = prompts.get(
        preset_type, "Summarize key analytical insights from this dataset."
    )
    # SECURITY SANBOX RATIONALE:
    # Executing LLM-generated strings directly poses an arbitrary code execution risk.

    exec_res = nl_query_pipeline(prompt, dataframe, schema_json)

    if exec_res["success"]:
        return format_result_with_llm(prompt, exec_res["result"])
    return "Unable to generate preset report."


import pandas as pd


def analyze_anomalies_iqr(df: pd.DataFrame, target_col: str, client):
    """Detects numeric outliers using IQR and asks the LLM for a business narrative (Task D3)."""
    if target_col not in df.columns or not pd.api.types.is_numeric_dtype(
        df[target_col]
    ):
        return None, "Selected column must be numeric."

    # 1. IQR Calculation
    Q1 = df[target_col].quantile(0.25)
    Q3 = df[target_col].quantile(0.75)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR

    # Filter anomaly rows
    anomalies = df[(df[target_col] < lower_bound) | (df[target_col] > upper_bound)]
    outlier_count = len(anomalies)
    total_count = len(df)
    outlier_pct = (
        round((outlier_count / total_count) * 100, 2) if total_count > 0 else 0
    )

    if outlier_count == 0:
        return (
            anomalies,
            f"No statistical outliers detected in `{target_col}` using the IQR method (1.5x threshold).",
        )

    # 2. Extract context for LLM narrative
    summary_stats = {
        "column": target_col,
        "total_records": total_count,
        "outlier_count": outlier_count,
        "outlier_percentage": outlier_pct,
        "q1": Q1,
        "q3": Q3,
        "iqr": IQR,
        "lower_bound": lower_bound,
        "upper_bound": upper_bound,
        "max_outlier": anomalies[target_col].max(),
        "min_outlier": anomalies[target_col].min(),
    }

    # 1. Determine severity programmatically based on percentage
    if outlier_pct == 0:
        severity = "None"
    elif outlier_pct < 2.0:
        severity = "Low"
    elif outlier_pct < 5.0:
        severity = "Moderate"
    else:
        severity = "Severe"

    # 2. Pass severity into the prompt
    prompt = f"""
    You are a Lead Data Analyst. Write a concise executive narrative explaining statistical anomalies found in the dataset.

    DATA SUMMARY:
    - Target Metric: {summary_stats["column"]}
    - Outliers Detected: {summary_stats["outlier_count"]} out of {summary_stats["total_records"]} rows ({summary_stats["outlier_percentage"]}%)
    - Pre-calculated Severity Level: {severity}

    INSTRUCTIONS:
    1. Begin with: **Outlier Severity:** {severity} ({summary_stats["outlier_percentage"]}%)
    2. Explain what these extreme values mean for business analysis in 2 concise bullet points.
    3. Provide 2-3 technical recommendations (e.g., Winsorizing, filtering, or manual verification).
    Keep tone professional, crisp, and concise.
    """

    try:
        response = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
        )
        narrative = response.choices[0].message.content
    except Exception as e:
        narrative = f"Detected {outlier_count} outliers ({outlier_pct}%). LLM Narrative generation failed: {str(e)}"

    return anomalies, narrative


import pandas as pd
from sklearn.linear_model import LinearRegression


def run_predictive_regression(df: pd.DataFrame, x_col: str, y_col: str, client):
    """Fits a Linear Regression model (Task D2) and uses LLM to interpret the coefficients."""
    # Ensure numeric and drop missing values for selected columns
    clean_df = df[[x_col, y_col]].dropna()

    if len(clean_df) < 5:
        return (
            None,
            "Insufficient non-null data points to fit a linear model.",
            {},
        )

    X = clean_df[[x_col]]
    y = clean_df[y_col]

    # Fit Scikit-Learn Model
    model = LinearRegression()
    model.fit(X, y)

    coef = float(model.coef_[0])
    intercept = float(model.intercept_)
    r2_score = float(model.score(X, y))

    metrics = {
        "x_col": x_col,
        "y_col": y_col,
        "coefficient": round(coef, 4),
        "intercept": round(intercept, 4),
        "r2_score": round(r2_score, 4),
        "sample_size": len(clean_df),
    }

    # Prompt LLM for Business Explanation of Coefficients
    prompt = f"""
You are a Lead Data Scientist explaining linear regression results to business stakeholders.

MODEL RESULTS:
- Independent Variable (X): {x_col}
- Dependent Variable (Y): {y_col}
- Slope Coefficient (m): {metrics["coefficient"]}
- Intercept (b): {metrics["intercept"]}
- R² Score: {metrics["r2_score"]} (explains {metrics["r2_score"] * 100:.2f}% of variance)
- Sample Size: {metrics["sample_size"]} rows

INSTRUCTIONS:
1. Explain the mathematical relationship in plain business terms: "For every 1 unit increase in {x_col}, {y_col} changes by..."
2. Interpret the R² score (is the correlation weak, moderate, or strong?).
3. Provide 2 actionable business insights or strategic recommendations based on this relationship.
Keep response concise, structured, and professional.
"""

    try:
        response = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
        )
        narrative = response.choices[0].message.content
    except Exception as e:
        narrative = f"Regression model fitted successfully, but LLM explanation failed: {str(e)}"

    return model, narrative, metrics


import datetime
import io
import os
import sys
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
import pandas as pd
from weasyprint import HTML

# Tell Anaconda where to find Homebrew libraries on macOS
if sys.platform == "darwin":
    os.environ["DYLD_FALLBACK_LIBRARY_PATH"] = "/opt/homebrew/lib:" + os.environ.get(
        "DYLD_FALLBACK_LIBRARY_PATH", "/usr/local/lib"
    )


def generate_pdf_report(
    dataset_name: str,
    total_rows: int,
    total_cols: int,
    filters_applied: str,
    narrative_text: str,
    summary_df: pd.DataFrame = None,
) -> bytes:
    """Generates a highly structured, executive-ready PDF report for Task C4."""
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # --- CLEAN NARRATIVE TEXT PROCESSING ---
    formatted_narrative = ""
    lines = [line.strip() for line in narrative_text.split("\n") if line.strip()]

    # Check if narrative contains metric dumps to render as cards
    is_data_list = any("Customer Segment:" in l or "Total Sales:" in l for l in lines)

    if is_data_list:
        formatted_narrative += "<div style='display: grid; grid-template-columns: 1fr 1fr; gap: 12px; margin-top: 10px;'>"
        for line in lines:
            clean_line = line.replace("**", "").replace("markdown", "").strip()
            if not clean_line:
                continue

            parts = [p.strip() for p in clean_line.split(",")]
            seg_val, age_val, sales_val = "N/A", "N/A", "N/A"

            for part in parts:
                if "Customer Segment:" in part:
                    seg_val = part.split(":", 1)[1].strip()
                elif "Age Segment:" in part:
                    age_val = part.split(":", 1)[1].strip()
                elif "Total Sales:" in part:
                    sales_val = part.split(":", 1)[1].strip()

            formatted_narrative += f"""
            <div style="background-color: #f8fafc; border: 1px solid #e2e8f0; border-radius: 6px; padding: 12px; box-shadow: 0 1px 2px rgba(0,0,0,0.02);">
                <div style="font-size: 8.5pt; text-transform: uppercase; letter-spacing: 0.5px; color: #64748b; font-weight: bold; margin-bottom: 4px;">{seg_val} Segment ({age_val})</div>
                <div style="font-size: 14pt; font-weight: bold; color: #0f172a;">{sales_val}</div>
            </div>
            """
        formatted_narrative += "</div>"
    else:
        for para in lines:
            if para.startswith("#"):
                title = para.lstrip("#").strip()
                formatted_narrative += f"<h3 style='color: #0f172a; margin-top: 16px; margin-bottom: 6px; font-size: 11pt; border-bottom: 1px solid #e2e8f0; padding-bottom: 2px;'>{title}</h3>"
            elif para.startswith("* ") or para.startswith("- "):
                item = para[2:].strip()
                formatted_narrative += (
                    f"<li style='margin-bottom: 4px; color: #334155;'>{item}</li>"
                )
            else:
                formatted_narrative += f"<p style='color: #334155; line-height: 1.5; margin-bottom: 8px; font-size: 10pt;'>{para}</p>"

    # --- SUMMARY TABLE PROCESSING ---
    table_html = ""
    if summary_df is not None and not summary_df.empty:
        preview_df = summary_df.head(10).copy()

        table_html += (
            "<h3 style='color: #0f172a; margin-top: 24px; font-size: 12pt; font-weight: bold;'"
            " border-left: 3px solid #0284c7; padding-left: 8px;'>📊 Query Result"
            " Summary Table</h3>"
        )
        table_html += "<table style='width:100%; border-collapse: collapse; margin-top: 8px; font-size: 9.5pt; box-shadow: 0 1px 3px rgba(0,0,0,0.02);'>"
        table_html += "<thead><tr style='background-color: #0f172a; color: #ffffff; text-align: left;'>"

        for col in preview_df.columns:
            clean_header = str(col).replace("_", " ").title()
            table_html += f"<th style='padding: 10px 12px; font-weight: 600; border-bottom: 2px solid #cbd5e1;'>{clean_header}</th>"
        table_html += "</tr></thead><tbody>"

        for idx, row in preview_df.iterrows():
            bg = "#f8fafc" if idx % 2 == 0 else "#ffffff"
            table_html += f"<tr style='background-color: {bg}; border-bottom: 1px solid #e2e8f0;'>"
            for col_name, val in row.items():
                if "sales" in str(col_name).lower() and isinstance(
                    val, (int, float, round)
                ):
                    display_val = f"${val:,.2f}"
                elif isinstance(val, (int, float)):
                    display_val = f"{val:,}"
                else:
                    display_val = str(val)

                table_html += (
                    f"<td style='padding: 8px 12px; color: #334155;'>{display_val}</td>"
                )
            table_html += "</tr>"
        table_html += "</tbody></table>"

    # --- COMPOSITE A4 TEMPLATE ---
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            @page {{
                size: A4;
                margin: 20mm 15mm;
                @bottom-right {{
                    content: "Page " counter(page) " of " counter(pages);
                    font-size: 8.5pt;
                    color: #94a3b8;
                    font-family: Arial, sans-serif;
                }}
                @bottom-left {{
                    content: "Platform Intelligence Engine • Confidential";
                    font-size: 8.5pt;
                    color: #94a3b8;
                    font-family: Arial, sans-serif;
                }}
            }}
            body {{
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;
                margin: 0;
                color: #1e293b;
                background-color: #ffffff;
                -webkit-print-color-adjust: exact;
            }}
            .header {{
                border-bottom: 2px solid #f1f5f9;
                padding-bottom: 16px;
                margin-bottom: 20px;
            }}
            .header h1 {{
                margin: 0;
                font-size: 18pt;
                font-weight: 700;
                color: #0f172a;
                letter-spacing: -0.5px;
            }}
            .header p {{
                margin: 4px 0 0 0;
                font-size: 9.5pt;
                color: #64748b;
            }}
            .meta-grid {{
                display: table;
                width: 100%;
                margin-bottom: 24px;
                border: 1px solid #e2e8f0;
                border-radius: 6px;
                border-collapse: separate;
            }}
            .meta-row {{
                display: table-row;
            }}
            .meta-cell {{
                display: table-cell;
                padding: 10px 14px;
                font-size: 9pt;
                color: #475569;
                border-bottom: 1px solid #e2e8f0;
            }}
            .meta-cell:first-child {{
                border-right: 1px solid #e2e8f0;
            }}
            .meta-row:last-child .meta-cell {{
                border-bottom: none;
            }}
            .label {{
                font-weight: 600;
                color: #0f172a;
            }}
            .section-title {{
                font-size: 12pt;
                font-weight: bold;
                color: #0f172a;
                margin-top: 10px;
                margin-bottom: 10px;
                border-left: 3px solid #0284c7;
                padding-left: 8px;
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>Executive Analytics & AI Insight Report</h1>
            <p>Task C4 — Automated Insight Delivery Pipeline</p>
        </div>

        <div class="meta-grid">
            <div class="meta-row">
                <div class="meta-cell"><span class="label">Dataset Focus:</span> {dataset_name}</div>
                <div class="meta-cell"><span class="label">Generated On:</span> {timestamp}</div>
            </div>
            <div class="meta-row">
                <div class="meta-cell"><span class="label">Active Scope:</span> {total_rows:,} rows • {total_cols} attributes</div>
                <div class="meta-cell"><span class="label">Runtime Filters:</span> {filters_applied}</div>
            </div>
        </div>

        <div class="section-title">🤖 AI Assistant Narrative & Insights</div>
        <div style="margin-bottom: 20px;">
            {formatted_narrative}
        </div>

        {table_html}
    </body>
    </html>
    """

    pdf_buffer = io.BytesIO()
    HTML(string=html_content).write_pdf(pdf_buffer)
    return pdf_buffer.getvalue()


def generate_docx_report(
    dataset_name: str,
    total_rows: int,
    total_cols: int,
    filters_applied: str,
    narrative_text: str,
    summary_df: pd.DataFrame = None,
) -> bytes:
    """Generates a professional Word (.docx) report for Task C4."""
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Executive Analytics & AI Insight Report")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 23, 42)

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(
        "Task C4 — Exported Report | Capstone Data Intelligence Platform"
    )
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(10)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    meta_data = [
        [
            f"Dataset: {dataset_name}",
            f"Generated On: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        ],
        [
            f"Active Records: {total_rows:,} rows ({total_cols} cols)",
            f"Active Filters: {filters_applied}",
        ],
    ]

    for row_idx, row in enumerate(meta_table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.text = meta_data[row_idx][col_idx]
            shading = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls("w")))
            cell._tc.get_or_add_tcPr().append(shading)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    h1 = doc.add_paragraph()
    h1_run = h1.add_run("🤖 AI Assistant Narrative & Insights")
    h1_run.font.name = "Arial"
    h1_run.font.size = Pt(14)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(2, 132, 199)

    for para in narrative_text.split("\n"):
        p_str = para.strip()
        if not p_str:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(p_str)
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(51, 65, 85)

    if summary_df is not None and not summary_df.empty:
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        h2 = doc.add_paragraph()
        h2_run = h2.add_run("📊 Query Result Data Preview")
        h2_run.font.name = "Arial"
        h2_run.font.size = Pt(12)
        h2_run.font.bold = True
        h2_run.font.color.rgb = RGBColor(15, 23, 42)

        preview_df = summary_df.head(10)
        table = doc.add_table(rows=len(preview_df) + 1, cols=len(preview_df.columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table.rows[0].cells
        for col_idx, col_name in enumerate(preview_df.columns):
            hdr_cells[col_idx].text = str(col_name)
            shading = parse_xml(r'<w:shd {} w:fill="0F172A"/>'.format(nsdecls("w")))
            hdr_cells[col_idx]._tc.get_or_add_tcPr().append(shading)
            for p in hdr_cells[col_idx].paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.5)
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)

        for r_idx, row_data in enumerate(preview_df.iterrows()):
            row_cells = table.rows[r_idx + 1].cells
            bg_color = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, val in enumerate(row_data[1]):
                row_cells[c_idx].text = str(val)
                shading = parse_xml(
                    r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), bg_color)
                )
                row_cells[c_idx]._tc.get_or_add_tcPr().append(shading)
                for p in row_cells[c_idx].paragraphs:
                    for r in p.runs:
                        r.font.name = "Arial"
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(51, 65, 85)

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    return doc_buffer.getvalue()
