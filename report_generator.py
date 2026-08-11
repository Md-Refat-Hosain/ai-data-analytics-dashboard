import datetime
import io
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
import pandas as pd
from weasyprint import HTML


def generate_pdf_report(
    dataset_name: str,
    total_rows: int,
    total_cols: int,
    filters_applied: str,
    narrative_text: str,
    summary_df: pd.DataFrame = None,
) -> bytes:
    """Generates a professional PDF report for Task C4 using HTML + WeasyPrint."""
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Convert narrative markdown to styled HTML
    formatted_narrative = ""
    for para in narrative_text.split("\n"):
        para_trimmed = para.strip()
        if not para_trimmed:
            continue
        if para_trimmed.startswith("#"):
            title = para_trimmed.lstrip("#").strip()
            formatted_narrative += f"<h3 style='color: #0f172a; margin-top: 15px; margin-bottom: 8px; border-bottom: 1px solid #e2e8f0; padding-bottom: 4px;'>{title}</h3>"
        elif para_trimmed.startswith("* ") or para_trimmed.startswith("- "):
            item = para_trimmed[2:].strip()
            formatted_narrative += (
                f"<li style='margin-bottom: 6px; color: #334155;'>{item}</li>"
            )
        else:
            formatted_narrative += f"<p style='color: #334155; line-height: 1.6; margin-bottom: 10px;'>{para_trimmed}</p>"

    # Optional DataFrame preview table
    table_html = ""
    if summary_df is not None and not summary_df.empty:
        preview_df = summary_df.head(10)
        table_html += "<h3 style='color: #0f172a; margin-top: 20px;'>📊 Query Result Summary Table</h3>"
        table_html += "<table style='width:100%; border-collapse: collapse; margin-top: 10px; font-size: 9.5pt;'>"
        table_html += "<thead><tr style='background-color: #0f172a; color: #ffffff; text-align: left;'>"
        for col in preview_df.columns:
            table_html += (
                f"<th style='padding: 8px; border: 1px solid #cbd5e1;'>{col}</th>"
            )
        table_html += "</tr></thead><tbody>"

        for idx, row in preview_df.iterrows():
            bg = "#f8fafc" if idx % 2 == 0 else "#ffffff"
            table_html += f"<tr style='background-color: {bg};'>"
            for val in row:
                table_html += f"<td style='padding: 6px 8px; border: 1px solid #e2e8f0; color: #334155;'>{val}</td>"
            table_html += "</tr>"
        table_html += "</tbody></table>"

    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            @page {{
                size: A4;
                margin: 20mm 15mm;
                background-color: #ffffff;
                @bottom-right {{
                    content: "Page " counter(page) " of " counter(pages);
                    font-size: 9pt;
                    color: #94a3b8;
                }}
                @bottom-left {{
                    content: "AI-Powered Data Platform Report";
                    font-size: 9pt;
                    color: #94a3b8;
                }}
            }}
            * {{ box-sizing: border-box; }}
            body {{
                font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
                margin: 0;
                padding: 0;
                color: #1e293b;
            }}
            .header-banner {{
                background-color: #0f172a;
                color: #ffffff;
                padding: 24px;
                border-radius: 8px;
                margin-bottom: 20px;
            }}
            .header-banner h1 {{
                margin: 0 0 6px 0;
                font-size: 20pt;
                color: #38bdf8;
            }}
            .header-banner p {{
                margin: 0;
                font-size: 10pt;
                color: #94a3b8;
            }}
            .meta-card {{
                background-color: #f1f5f9;
                border-left: 4px solid #0284c7;
                padding: 12px 16px;
                margin-bottom: 20px;
                border-radius: 0 6px 6px 0;
            }}
            .meta-grid {{
                display: table;
                width: 100%;
            }}
            .meta-cell {{
                display: table-cell;
                width: 50%;
                font-size: 10pt;
                color: #334155;
            }}
            .meta-label {{
                font-weight: bold;
                color: #0f172a;
            }}
            .section-title {{
                font-size: 14pt;
                color: #0f172a;
                border-bottom: 2px solid #0284c7;
                padding-bottom: 4px;
                margin-top: 20px;
                margin-bottom: 12px;
            }}
        </style>
    </head>
    <body>
        <div class="header-banner">
            <h1>Executive Analytics & AI Insight Report</h1>
            <p>Task C4 — Exported Report | Capstone Data Intelligence Platform</p>
        </div>

        <div class="meta-card">
            <div class="meta-grid">
                <div class="meta-cell">
                    <p><span class="meta-label">Dataset:</span> {dataset_name}</p>
                    <p><span class="meta-label">Active Records:</span> {total_rows:,} rows | {total_cols} columns</p>
                </div>
                <div class="meta-cell">
                    <p><span class="meta-label">Generated On:</span> {timestamp}</p>
                    <p><span class="meta-label">Active Filters:</span> {filters_applied}</p>
                </div>
            </div>
        </div>

        <div class="section-title">🤖 AI Assistant Narrative & Insights</div>
        <div>
            {formatted_narrative}
        </div>

        {table_html}
    </body>
    </html>
    """

    pdf_buffer = io.BytesIO()
    HTML(string=html_content).write_pdf(pdf_buffer)
    return pdf_buffer.getvalue()


def generate_docx_report(
    dataset_name: str,
    total_rows: int,
    total_cols: int,
    filters_applied: str,
    narrative_text: str,
    summary_df: pd.DataFrame = None,
) -> bytes:
    """Generates a professional Word (.docx) report for Task C4."""
    doc = docx.Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Document Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Executive Analytics & AI Insight Report")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 23, 42)

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(
        "Task C4 — Exported Report | Capstone Data Intelligence Platform"
    )
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(10)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Metadata Box Table
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    meta_data = [
        [
            f"Dataset: {dataset_name}",
            f"Generated On: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        ],
        [
            f"Active Records: {total_rows:,} rows ({total_cols} cols)",
            f"Active Filters: {filters_applied}",
        ],
    ]

    for row_idx, row in enumerate(meta_table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.text = meta_data[row_idx][col_idx]
            shading = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls("w")))
            cell._tc.get_or_add_tcPr().append(shading)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Narrative Section
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("🤖 AI Assistant Narrative & Insights")
    h1_run.font.name = "Arial"
    h1_run.font.size = Pt(14)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(2, 132, 199)

    for para in narrative_text.split("\n"):
        p_str = para.strip()
        if not p_str:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(p_str)
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(51, 65, 85)

    # DataFrame Preview
    if summary_df is not None and not summary_df.empty:
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        h2 = doc.add_paragraph()
        h2_run = h2.add_run("📊 Query Result Data Preview")
        h2_run.font.name = "Arial"
        h2_run.font.size = Pt(12)
        h2_run.font.bold = True
        h2_run.font.color.rgb = RGBColor(15, 23, 42)

        preview_df = summary_df.head(10)
        table = doc.add_table(rows=len(preview_df) + 1, cols=len(preview_df.columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table.rows[0].cells
        for col_idx, col_name in enumerate(preview_df.columns):
            hdr_cells[col_idx].text = str(col_name)
            shading = parse_xml(r'<w:shd {} w:fill="0F172A"/>'.format(nsdecls("w")))
            hdr_cells[col_idx]._tc.get_or_add_tcPr().append(shading)
            for p in hdr_cells[col_idx].paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9.5)
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)

        for r_idx, row_data in enumerate(preview_df.iterrows()):
            row_cells = table.rows[r_idx + 1].cells
            bg_color = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, val in enumerate(row_data[1]):
                row_cells[c_idx].text = str(val)
                shading = parse_xml(
                    r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), bg_color)
                )
                row_cells[c_idx]._tc.get_or_add_tcPr().append(shading)
                for p in row_cells[c_idx].paragraphs:
                    for r in p.runs:
                        r.font.name = "Arial"
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(51, 65, 85)

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    return doc_buffer.getvalue()
